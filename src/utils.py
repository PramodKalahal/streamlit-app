import pandas as pd
import numpy as np
import re
import os
from collections import Counter
import random
import pdfplumber
import docx

def GET_MASTER_SKILLS():
    """Returns a sorted list of all supported technical skills."""
    return sorted([
        'python', 'java', 'sql', 'c++', 'c#', 'aws', 'azure', 'gcp', 'spark', 
        'pandas', 'numpy', 'tensorflow', 'pytorch', 'excel', 'tableau', 
        'power bi', 'react', 'node.js', 'javascript', 'html', 'css', 'html/css',
        'machine learning', 'data science', 'deep learning', 'docker', 'kubernetes',
        'mongodb', 'postgresql', 'flask', 'django', 'spring boot', 'scala',
        'typescript', 'tailwind css', 'redux', 'rest api', 'algorithms', 'data structures',
        'system design', 'linux', 'scikit-learn', 'feature engineering', 'jupyter',
        'data visualization', 'neural networks', 'mlops', 'aws sagemaker', 'model deployment',
        'nlp', 'google data studio', 'metabase', 'data cleaning', 'reporting',
        'fastapi', 'pytest', 'celery', 'redis', 'express', 'next.js', 'git',
        'jenkins', 'terraform', 'ci/cd', 'ansible', 'bash', 'figma', 'adobe xd',
        'user research', 'wireframing', 'color theory', 'prototyping', 'responsive design',
        'iam', 'vpc', 'serverless', 'microservices', 'cloud migration', 'hadoop',
        'kafka', 'airflow', 'etl', 'snowflake', 'dbt', 'penetration testing',
        'siem', 'ids/ips', 'cryptography', 'vulnerability assessment', 'firewalls', 'security audit'
    ])

def extract_text(uploaded_file):
    """Extracts text from PDF or Word documents."""
    text = ""
    try:
        if uploaded_file.name.lower().endswith('.pdf'):
            with pdfplumber.open(uploaded_file) as pdf:
                for page in pdf.pages:
                    content = page.extract_text()
                    if content:
                        text += content + "\n"
        elif uploaded_file.name.lower().endswith('.docx') or uploaded_file.name.lower().endswith('.doc'):
            try:
                doc = docx.Document(uploaded_file)
                # Extract text from paragraphs
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                # Extract text from tables (common in resumes)
                table_text = []
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                table_text.append(cell.text.strip())
                
                text = "\n".join(paragraphs + table_text)
            except Exception as doc_err:
                # Fallback if docx.Document fails on old .doc files
                print(f"Docx error: {doc_err}")
                text = str(uploaded_file.read())
    except Exception as e:
        print(f"General extraction error: {e}")
    
    return text.strip()

def extract_objective_section(text):
    """
    Extracts the objective or summary section from the resume text.
    Looks for headers like 'Objective', 'Summary', 'Profile', etc.
    """
    headers = ['objective', 'summary', 'professional summary', 'profile', 'about me', 'career goal']
    lines = text.split('\n')
    objective_lines = []
    found = False
    
    for i, line in enumerate(lines):
        clean_line = line.strip().lower()
        if any(h in clean_line for h in headers) and len(clean_line) < 30:
            found = True
            continue
        
        if found:
            # If we hit another likely header, stop
            if len(line.strip()) > 0 and line.strip().isupper() and len(line.strip()) < 30:
                break
            if i + 1 < len(lines) and len(lines[i+1].strip()) > 0 and lines[i+1].strip().isupper():
                objective_lines.append(line)
                break
            objective_lines.append(line)
            
            # Limit objective to 10 lines
            if len(objective_lines) > 10:
                break
                
    if found and objective_lines:
        return "\n".join(objective_lines).strip()
    return text[:500] # Fallback to first 500 chars

def determine_target_role(text):
    """
    Analyzes resume text to determine the most likely target role.
    Uses three-level weighted scoring as per documentation:
    - Primary Keywords: 10 pts
    - Technical Skills: 2 pts
    - Contextual Keywords: 1 pt
    """
    text = text.lower()
    scores = []
    
    for role, data in ROLE_PATTERNS.items():
        score = 0
        
        # Level 1: Primary Keywords (10 pts each)
        for pk in data.get('primary', []):
            if pk.lower() in text:
                score += 10
                
        # Level 2: Technical Skills (2 pts each)
        for skill in data.get('skills', []):
            if skill.lower() in text:
                score += 2
                
        # Level 3: Contextual Keywords (1 pt each)
        for ctx in data.get('context', []):
            if ctx.lower() in text:
                score += 1
                
        scores.append({'role': role, 'score': score})
    
    if not scores:
        return {'target_role': "Unknown", 'confidence': "Low", 'reason': "No roles matched.", 'score': 0}

    sorted_scores = sorted(scores, key=lambda x: x['score'], reverse=True)
    top_role = sorted_scores[0]
    runner_up = sorted_scores[1] if len(sorted_scores) > 1 else {'score': 0}
    
    score_gap = top_role['score'] - runner_up['score']
    
    if top_role['score'] >= 4: # Minimum threshold
        if top_role['score'] >= 15 and score_gap >= 5:
            confidence = "High"
        elif top_role['score'] >= 8 and score_gap >= 3:
            confidence = "Medium"
        else:
            confidence = "Low"
            
        return {
            'target_role': top_role['role'],
            'confidence': confidence,
            'reason': f"Detected alignment with {top_role['role']} (Score: {top_role['score']}, Gap: {score_gap}).",
            'score': top_role['score']
        }
    
    return {
        'target_role': "Unknown",
        'confidence': "Low",
        'reason': "Could not find a strong match for a specific role in the provided text.",
        'score': top_role['score']
    }

def get_all_suitable_roles(text):
    """Returns a list of roles that match the resume text."""
    text = text.lower()
    suitable_roles = []
    
    for role, data in ROLE_PATTERNS.items():
        # Identify which specific skills matched
        matched = [skill for skill in data['skills'] if skill.lower() in text]
        if len(matched) >= 2:
            suitable_roles.append({
                'role': role, 
                'score': len(matched),
                'matched_skills': matched
            })
            
    return sorted(suitable_roles, key=lambda x: x['score'], reverse=True)

def calculate_ats_score(text, matched, missing):
    """Calculates an ATS compatibility score based on skills and formatting."""
    score = 0
    breakdown = {}
    feedback = []
    
    # 1. Skill Match (Max 60 points)
    skill_match_ratio = len(matched) / (len(matched) + len(missing)) if (len(matched) + len(missing)) > 0 else 0
    skill_points = int(skill_match_ratio * 60)
    score += skill_points
    breakdown['Skill Alignment'] = skill_points
    
    if skill_match_ratio < 0.5:
        feedback.append("Your resume is missing over 50% of the core skills required for this role.")
    
    # 2. Content Length (Max 20 points)
    word_count = len(text.split())
    length_points = min(20, (word_count // 100) * 5)
    score += length_points
    breakdown['Content Depth'] = length_points
    
    if word_count < 300:
        feedback.append("The resume content seems a bit thin. Try adding more detailed project descriptions.")
    
    # 3. Formatting & Structure (Max 20 points)
    format_points = 0
    if re.search(r'experience|work history', text.lower()): format_points += 7
    if re.search(r'education', text.lower()): format_points += 7
    if re.search(r'projects|portfolio', text.lower()): format_points += 6
    score += format_points
    breakdown['Structural Integrity'] = format_points
    
    if format_points < 20:
        feedback.append("Some standard sections (Experience, Education, Projects) appear to be missing or poorly labeled.")

    return {
        'score': min(100, score),
        'breakdown': breakdown,
        'feedback': feedback
    }

# --- CONSTANTS & MOCK DATA ---
ROLE_PATTERNS = {
    'Web Developer': {
        'primary': ['web developer', 'frontend developer', 'ui developer'],
        'skills': ['JavaScript', 'React', 'Node.js', 'HTML/CSS', 'TypeScript', 'Tailwind CSS', 'Redux', 'REST API'],
        'context': ['ui', 'ux', 'responsive', 'frontend', 'browser', 'web'],
        'base_salary': 700000
    },
    'Software Engineer': {
        'primary': ['software engineer', 'software developer', 'backend engineer'],
        'skills': ['Java', 'C++', 'Python', 'C#', 'Algorithms', 'Data Structures', 'System Design', 'Linux'],
        'context': ['development', 'engineering', 'scalable', 'backend', 'system'],
        'base_salary': 900000
    },
    'Data Scientist': {
        'primary': ['data scientist', 'statistical modeler', 'ai researcher'],
        'skills': ['Python', 'Statistics', 'Pandas', 'NumPy', 'Scikit-Learn', 'Feature Engineering', 'Jupyter', 'Data Visualization'],
        'context': ['insights', 'research', 'experiment', 'quantitative', 'statistical'],
        'base_salary': 900000
    },
    'Machine Learning Engineer': {
        'primary': ['ml engineer', 'machine learning engineer', 'mlops engineer'],
        'skills': ['Python', 'TensorFlow', 'PyTorch', 'Neural Networks', 'MLOps', 'AWS Sagemaker', 'Model Deployment', 'NLP'],
        'context': ['production', 'scalable', 'pipeline', 'infrastructure', 'model'],
        'base_salary': 1000000
    },
    'Data Analyst': {
        'primary': ['data analyst', 'bi analyst', 'business intelligence analyst'],
        'skills': ['SQL', 'Excel', 'Tableau', 'Power BI', 'Google Data Studio', 'Metabase', 'Data Cleaning', 'Reporting'],
        'context': ['report', 'visualization', 'insights', 'dashboard', 'trends'],
        'base_salary': 600000
    },
    'Python Developer': {
        'primary': ['python developer', 'python engineer', 'django developer'],
        'skills': ['Python', 'Django', 'Flask', 'FastAPI', 'Pytest', 'Celery', 'PostgreSQL', 'Redis'],
        'context': ['backend', 'scripting', 'api', 'server', 'asynchronous'],
        'base_salary': 750000
    },
    'Full Stack Developer': {
        'primary': ['full stack developer', 'fullstack engineer', 'mern developer'],
        'skills': ['React', 'Node.js', 'Express', 'MongoDB', 'Next.js', 'PostgreSQL', 'Docker', 'Git'],
        'context': ['frontend', 'backend', 'full-stack', 'deployment', 'end-to-end'],
        'base_salary': 850000
    },
    'DevOps Engineer': {
        'primary': ['devops engineer', 'site reliability engineer', 'sre'],
        'skills': ['Docker', 'Kubernetes', 'Jenkins', 'Terraform', 'CI/CD', 'Ansible', 'Bash', 'AWS'],
        'context': ['automation', 'infrastructure', 'cloud', 'deployment', 'monitoring'],
        'base_salary': 950000
    },
    'UI/UX Designer': {
        'primary': ['ui/ux designer', 'product designer', 'ux researcher'],
        'skills': ['Figma', 'Adobe XD', 'User Research', 'Wireframing', 'Color Theory', 'Prototyping', 'Responsive Design'],
        'context': ['design', 'user', 'interface', 'experience', 'visual'],
        'base_salary': 650000
    },
    'Cloud Architect': {
        'primary': ['cloud architect', 'solutions architect', 'cloud engineer'],
        'skills': ['AWS', 'Azure', 'GCP', 'IAM', 'VPC', 'Serverless', 'Microservices', 'Cloud Migration'],
        'context': ['cloud', 'architecture', 'scalability', 'migration', 'hybrid'],
        'base_salary': 1200000
    },
    'Data Engineer': {
        'primary': ['data engineer', 'elt engineer', 'big data engineer'],
        'skills': ['Spark', 'Hadoop', 'Kafka', 'Airflow', 'ETL', 'Snowflake', 'Dbt', 'SQL'],
        'context': ['pipeline', 'warehouse', 'processing', 'lake', 'streaming'],
        'base_salary': 950000
    },
    'Cybersecurity Analyst': {
        'primary': ['cybersecurity analyst', 'security engineer', 'soc analyst'],
        'skills': ['Penetration Testing', 'SIEM', 'IDS/IPS', 'Cryptography', 'Vulnerability Assessment', 'Firewalls', 'Security Audit'],
        'context': ['security', 'threat', 'protection', 'network', 'compliance'],
        'base_salary': 850000
    },
    'Business Analyst': {
        'primary': ['business analyst', 'systems analyst', 'product analyst'],
        'skills': ['SQL', 'Excel', 'Tableau', 'Requirement Gathering', 'Process Mapping', 'Agile', 'Jira'],
        'context': ['requirements', 'business', 'process', 'stakeholders', 'specs'],
        'base_salary': 700000
    },
    'AI Engineer': {
        'primary': ['ai engineer', 'artificial intelligence engineer', 'generative ai engineer'],
        'skills': ['Python', 'PyTorch', 'TensorFlow', 'LLMs', 'OpenCV', 'Reinforcement Learning', 'NLP'],
        'context': ['neural', 'intelligence', 'autonomous', 'generative', 'transformers'],
        'base_salary': 1100000
    },
    'NLP Specialist': {
        'primary': ['nlp specialist', 'nlp engineer', 'natural language processing engineer'],
        'skills': ['Python', 'NLTK', 'SpaCy', 'Transformers', 'BERT', 'Tokenization', 'Word Embeddings'],
        'context': ['language', 'text', 'linguistics', 'nlp', 'embedding'],
        'base_salary': 1000000
    }
}

# --- INTERVIEW QUESTION BANK WITH ANSWERS (110+ Questions) ---
INTERVIEW_BANK = {
    "Python": [
        {"q": "How does memory management work in Python?", "a": "Python uses a private heap to manage memory. It has a built-in garbage collector that uses reference counting and a cyclic garbage collector to reclaim memory from objects that are no longer in use."},
        {"q": "What are decorators in Python?", "a": "Decorators are a way to modify or enhance functions or classes without changing their source code. They are commonly used for logging, access control, and memoization."},
        {"q": "Explain the Global Interpreter Lock (GIL).", "a": "The GIL is a mutex that protects access to Python objects, preventing multiple threads from executing bytecodes at once."},
        {"q": "Difference between Lists and Tuples?", "a": "Lists are mutable and use []. Tuples are immutable and use (). Tuples are generally faster and safer for constants."},
        {"q": "What is Python's 'self' keyword?", "a": "Self represents the instance of the class. It allows access to the attributes and methods of the class."},
        {"q": "What are lambda functions?", "a": "Small anonymous functions defined with the lambda keyword. They can have any number of arguments but only one expression."},
        {"q": "How do you manage dependencies in Python?", "a": "Typically using pip with requirements.txt, or environment managers like venv, conda, or poetry."},
        {"q": "What is the difference between deep and shallow copy?", "a": "Shallow copy creates a new object but references the same nested objects. Deep copy creates a new object and recursively copies all nested objects."},
        {"q": "What are generators in Python?", "a": "Functions that return an iterator using the 'yield' keyword, allowing for efficient processing of large datasets without loading everything into memory."},
        {"q": "Explain *args and **kwargs.", "a": "*args allows a function to accept any number of positional arguments. **kwargs allows for any number of keyword arguments."}
    ],
    "SQL": [
        {"q": "Difference between JOIN and Subquery?", "a": "JOINs combine data from multiple tables by linking rows. Subqueries are nested queries. JOINs are generally more performant for large datasets."},
        {"q": "What are Indexes and their types?", "a": "Indexes speed up data retrieval. Types include Clustered (determines physical storage order) and Non-Clustered (separate structure)."},
        {"q": "Explain ACID properties.", "a": "Atomicity, Consistency, Isolation, and Durability - the golden standard for reliable transactions."},
        {"q": "Difference between TRUNCATE and DELETE?", "a": "DELETE removes rows one by one and can be rolled back. TRUNCATE resets the table and is much faster but harder to roll back."},
        {"q": "What is a Stored Procedure?", "a": "A collection of SQL statements saved in the database that can be executed as a single unit or function."},
        {"q": "Explain normalization.", "a": "The process of organizing data to reduce redundancy (1NF, 2NF, 3NF)."},
        {"q": "What is a Window Function?", "a": "A function that performs calculations across a set of table rows that are somehow related to the current row (e.g., OVER, RANK)."},
        {"q": "Explain the difference between WHERE and HAVING.", "a": "WHERE filters rows before aggregation; HAVING filters groups after the GROUP BY clause."},
        {"q": "What is a View in SQL?", "a": "A virtual table based on the result-set of an SQL statement."},
        {"q": "What is a primary key vs a foreign key?", "a": "Primary key uniquely identifies a record in its own table. Foreign key links to a primary key in another table to establish a relationship."}
    ],
    "Java": [
        {"q": "Explain the Four Pillars of OOP.", "a": "Encapsulation, Abstraction, Inheritance, and Polymorphism."},
        {"q": "JDK vs JRE vs JVM?", "a": "JVM runs bytecode. JRE includes JVM and libraries. JDK includes JRE plus development tools like the compiler."},
        {"q": "Interface vs Abstract class?", "a": "Interfaces define behavior (what to do); Abstract classes define identity (what it is) and can store common state."},
        {"q": "What is the 'static' keyword?", "a": "Static members belong to the class itself rather than instances of the class."},
        {"q": "What is garbage collection?", "a": "Java's automatic process of reclaiming memory by deleting objects that are no longer reachable."},
        {"q": "Explain the Final, Finally, and Finalize keywords.", "a": "Final defines constants or prevents overriding. Finally is a block for cleanup. Finalize is a method called before object destruction (deprecated)."},
        {"q": "What is Polymorphism?", "a": "The ability of a single action to perform differently based on the object (Method Overloading and Method Overriding)."},
        {"q": "What is a thread in Java?", "a": "A lightweight sub-process, the smallest unit of processing. Java supports multithreading for concurrent execution."}
    ],
    "React & Frontend": [
        {"q": "What is the Virtual DOM?", "a": "A lightweight representation of the real DOM. React updates the Virtual DOM first, calculates differences, and only updates necessary parts of the real DOM."},
        {"q": "State vs Props?", "a": "State is internal and managed by the component. Props are external and passed down from parents (immutable to the receiver)."},
        {"q": "What are React Hooks?", "a": "Functions that let you use state and lifecycle features in functional components (e.g., useState, useEffect)."},
        {"q": "What is JSX?", "a": "A syntax extension that looks like HTML but is converted into standard JavaScript calls."},
        {"q": "What is Redux?", "a": "A state management library that provides a single source of truth for the entire application state."},
        {"q": "Explain Closure in JavaScript.", "a": "A function that remembers and accesses variables from its lexical scope even when executed outside that scope."},
        {"q": "What is a CSS Preprocessor?", "a": "A tool like SASS or LESS that extends CSS with features like variables, nesting, and mixins."},
        {"q": "Explain Responsive Design.", "a": "An approach to web design where pages render well on a variety of devices and window/screen sizes using media queries."},
        {"q": "What is an event loop in JS?", "a": "The mechanism that allows JS to be non-blocking by offloading operations to the system and handling their results when free."},
        {"q": "What is the difference between '==' and '==='?", "a": "'==' performs type coercion before comparison. '===' compares both value and type without coercion."}
    ],
    "Node.js & Backend": [
        {"q": "What is Node.js?", "a": "A runtime environment built on Chrome's V8 engine for executing JavaScript on the server side."},
        {"q": "What is Middleware in Express?", "a": "Functions that have access to the request/response objects and execute in the middle of processing a request."},
        {"q": "REST vs GraphQL?", "a": "REST uses standard HTTP methods for fixed resources. GraphQL allows clients to request exactly the data they need in one call."},
        {"q": "What is an API Gateway?", "a": "A server that acts as an entry point for microservices, handling routing, security, and rate limiting."},
        {"q": "Explain JWT (JSON Web Token).", "a": "A compact, URL-safe means of representing claims to be transferred between two parties, often used for authentication."},
        {"q": "What is Caching (e.g., Redis)?", "a": "Storing frequently accessed data in memory to speed up retrieval and reduce database load."},
        {"q": "Explain a Microservices architecture.", "a": "Building an application as a collection of small, independent services that communicate over a network."},
        {"q": "What is Spring Boot?", "a": "A Java-based framework used to create stand-alone, production-grade Spring-based applications easily."},
        {"q": "What is Django?", "a": "A high-level Python web framework that encourages rapid development and clean, pragmatic design."},
        {"q": "Explain the concept of WebSockets.", "a": "A protocol providing full-duplex communication channels over a single TCP connection, ideal for real-time apps."}
    ],
    "Machine Learning & AI": [
        {"q": "Supervised vs Unsupervised learning?", "a": "Supervised uses labeled data (known answers). Unsupervised finds patterns in unlabeled data."},
        {"q": "What is Overfitting?", "a": "When a model learns training data noise, failing to generalize to new data. Solved via regularization or more data."},
        {"q": "Explain Bias vs Variance.", "a": "Bias is error from wrong assumptions. Variance is error from sensitivity to small fluctuations in training data."},
        {"q": "What is Cross-Validation?", "a": "Splitting data into multiple sets to train and test the model multiple times for more reliable performance metrics."},
        {"q": "Explain a Neural Network.", "a": "A series of algorithms that mimic the human brain to recognize relationships in data through layers of neurons."},
        {"q": "What is NLP (Natural Language Processing)?", "a": "A field of AI focused on the interaction between computers and human language."},
        {"q": "What is Gradient Descent?", "a": "An optimization algorithm used to minimize a function (loss function) by iteratively moving in the direction of steepest descent."},
        {"q": "What is a Random Forest?", "a": "An ensemble of decision trees that improves classification accuracy and controls overfitting."},
        {"q": "Explain Precision and Recall.", "a": "Precision: Accuracy of positive predictions. Recall: Ability to find all actual positive instances."},
        {"q": "What is Reinforcement Learning?", "a": "AI training based on rewarding desired behaviors and punishing undesired ones."}
    ],
    "Data Science & Big Data": [
        {"q": "What is Pandas?", "a": "A library providing high-performance, easy-to-use data structures and data analysis tools for Python."},
        {"q": "Explain Spark's RDD.", "a": "Resilient Distributed Dataset - a fundamental data structure of Spark that is fault-tolerant and distributed."},
        {"q": "What is Hadoop?", "a": "A framework for distributed storage and processing of massive datasets using HDFS and MapReduce."},
        {"q": "What is an ETL process?", "a": "Extract, Transform, and Load - the process of moving data from source systems to a data warehouse."},
        {"q": "Explain Data Wrangling.", "a": "The process of cleaning, structuring, and enriching raw data into a desired format for better decision making."},
        {"q": "What is a Data Warehouse?", "a": "A central repository of integrated data from one or more disparate sources for reporting and analysis."},
        {"q": "Explain the 'p-value' in statistics.", "a": "The probability of obtaining results at least as extreme as the observed results, assuming the null hypothesis is true."},
        {"q": "What is Snowflake?", "a": "A cloud-based data platform provided as a Software-as-a-Service (SaaS)."},
        {"q": "What is Apache Kafka?", "a": "A distributed streaming platform used for building real-time data pipelines and streaming apps."},
        {"q": "Explain A/B Testing.", "a": "A statistical way to compare two versions (A and B) to determine which performs better for a specific metric."}
    ],
    "AWS & Cloud": [
        {"q": "What is AWS S3?", "a": "Simple Storage Service: Scalable object storage for any amount of data."},
        {"q": "Explain AWS EC2.", "a": "Elastic Compute Cloud: Virtual servers in the cloud."},
        {"q": "What is a VPC?", "a": "Virtual Private Cloud: A logically isolated section of the AWS cloud."},
        {"q": "What is Serverless computing?", "a": "Running code without managing servers (e.g., AWS Lambda). You only pay for execution time."},
        {"q": "IAM Policy vs Role?", "a": "Policies define permissions. Roles are identities that can be assumed by users or services to gain those permissions."},
        {"q": "What is CloudWatch?", "a": "A monitoring and observability service for AWS resources and applications."},
        {"q": "Explain Azure Resource Groups.", "a": "A container that holds related resources for an Azure solution."},
        {"q": "What is GCP BigQuery?", "a": "A fully managed, serverless data warehouse that enables scalable analysis over petabytes of data."},
        {"q": "What is Multi-Cloud strategy?", "a": "Using cloud services from multiple providers (AWS, Azure, GCP) to avoid vendor lock-in and increase resilience."},
        {"q": "Explain Edge Computing.", "a": "Distributed computing that brings computation and data storage closer to the location where it is needed."}
    ],
    "DevOps & Tools": [
        {"q": "What is Docker?", "a": "A platform for developers to develop, ship, and run applications inside containers."},
        {"q": "Kubernetes Pod vs Service?", "a": "A Pod is the smallest deployable unit (containers). A Service is an abstraction to expose those pods as a network service."},
        {"q": "What is Terraform?", "a": "An Infrastructure as Code (IaC) tool that allows you to build, change, and version infrastructure safely."},
        {"q": "Explain Git Rebase vs Merge.", "a": "Merge combines histories, potentially creating a merge commit. Rebase rewrites history to be linear."},
        {"q": "What is a CI/CD Pipeline?", "a": "Automation that handles building, testing, and deploying code consistently."},
        {"q": "What is Ansible?", "a": "An automation tool used for configuration management, application deployment, and task automation."},
        {"q": "Explain the 'Log' command in Git.", "a": "Shows the commit history of the current branch."},
        {"q": "What is Prometheus?", "a": "An open-source monitoring and alerting toolkit focused on time-series data."},
        {"q": "Explain Blue-Green Deployment.", "a": "A strategy that uses two identical environments (Blue is live, Green is new) to minimize downtime during updates."},
        {"q": "What is Selenium?", "a": "A suite of tools for automating web browser testing."}
    ],
    "Cybersecurity": [
        {"q": "What is the CIA triad?", "a": "Confidentiality, Integrity, and Availability - the core pillars of security."},
        {"q": "What is Phishing?", "a": "Social engineering where attackers masquerade as trusted entities to steal sensitive data."},
        {"q": "What is a Firewall?", "a": "A security system that controls incoming and outgoing network traffic based on rules."},
        {"q": "Explain Encryption (Symmetric vs Asymmetric).", "a": "Symmetric uses one key for both. Asymmetric uses a public key to encrypt and a private key to decrypt."},
        {"q": "What is a Penetration Test?", "a": "A simulated cyber attack against your computer system to check for exploitable vulnerabilities."},
        {"q": "What is a DDoS attack?", "a": "Distributed Denial of Service - overwhelming a system with traffic from multiple sources to make it unavailable."},
        {"q": "Explain Multi-Factor Authentication (MFA).", "a": "Requiring multiple forms of verification (password + code) to grant access."},
        {"q": "What is Hashing?", "a": "Converting data into a fixed-length string of characters (one-way process), commonly used for passwords."}
    ],
    "UI/UX Design": [
        {"q": "What is the difference between UI and UX?", "a": "UI is the visual aspect (look); UX is the overall experience and usability (feel)."},
        {"q": "What is Figma?", "a": "A collaborative web-based design and prototyping tool."},
        {"q": "Explain User Persona.", "a": "A fictional character created to represent a user type that might use your product."},
        {"q": "What is Wireframing?", "a": "A blueprint of a user interface that outlines structure and layout without visual details."},
        {"q": "What is Information Architecture?", "a": "Organizing and labeling websites and apps to support findability and usability."},
        {"q": "Explain Accessibility in Design.", "a": "Designing products so that people with disabilities can use them successfully."}
    ],
    "D.Viz & Analytics": [
        {"q": "Tableau vs Power BI?", "a": "Tableau is powerful for complex visuals; Power BI is integrated with MS ecosystem and easier for basic reporting."},
        {"q": "What is a Dashboard?", "a": "A visual display of the most important information needed to achieve one or more objectives."},
        {"q": "What are KPIs?", "a": "Key Performance Indicators - measurable values that demonstrate how effectively a company is achieving objectives."},
        {"q": "Difference between Dimension and Measure?", "a": "Dimensions are qualitative data (e.g., City). Measures are quantitative/numerical values (e.g., Sales)."},
        {"q": "What is a Pivot Table?", "a": "A data processing tool used to summarize, sort, and reorganize data in a spreadsheet or database."},
        {"q": "Explain Data Storytelling.", "a": "Using data, visuals, and narrative to communicate insights to a specific audience."},
        {"q": "What is Google Data Studio (Looker Studio)?", "a": "A free tool for converting data into informative, easy to read, easy to share, and fully customizable dashboards."},
        {"q": "What is an Outlier?", "a": "A data point that differs significantly from other observations."}
    ],
    "Soft Skills & Leadership": [
        {"q": "How do you explain technical concepts to non-technical stakeholders?", "a": "Focus on the 'why' and the 'impact' rather than the 'how'. Use analogies and avoid jargon."},
        {"q": "Tell me about a time you disagreed with a team lead. How did you handle it?", "a": "Explain how you presented your data-backed viewpoint respectfully and reached a consensus or followed the lead's decision professionally."},
        {"q": "What is your process for learning a new technology quickly?", "a": "Mention building a small proof-of-concept, reading official documentation, and following best practices from the community."},
        {"q": "Describe a situation where you had to handle a high-pressure deadline.", "a": "Focus on prioritization, clear communication with the team, and maintaining code quality under pressure."},
        {"q": "How do you ensure your work aligns with the company's business goals?", "a": "By understanding the KPIs, asking clarify questions about project impact, and staying focused on delivering value."}
    ],
    "Behavioral (Interviewer)": [
        {"q": "Tell me about yourself.", "a": "Highlight past experience, current role, and future goals relative to the company."},
        {"q": "What is your greatest weakness?", "a": "Mention a real skill gap you've worked on overcoming."},
        {"q": "Why should we hire you?", "a": "Focus on how your skills solve the company's specific problems."},
        {"q": "Describe a conflict you had with a co-worker.", "a": "Focus on how you resolved it professionally."},
        {"q": "Where do you see yourself in 5 years?", "a": "Align your ambition with the company's growth path."},
        {"q": "What is your proudest accomplishment?", "a": "Share a specific measurable result from a project."},
        {"q": "How do you handle high-pressure situations?", "a": "Talk about prioritization and staying calm/organized."},
        {"q": "Why are you leaving your current job?", "a": "Stay positive and focus on seeking new challenges/growth."},
        {"q": "What do you know about our company?", "a": "Show you've researched their mission, products, and culture."},
        {"q": "Do you have any questions for me?", "a": "Asking questions shows interest. Ask about team culture or project goals."}
    ]
}

def get_job_recommendations(user_skills):
    """Suggests roles based on matching skills."""
    role_scores = []
    user_skills_set = set([s.lower() for s in user_skills])
    
    for role, data in ROLE_PATTERNS.items():
        role_skills_set = set([s.lower() for s in data['skills']])
        match_count = len(user_skills_set & role_skills_set)
        match_score = (match_count / len(role_skills_set)) * 100
        role_scores.append({
            'role': role,
            'match_score': match_score,
            'skills': data['skills'],
            'avg_salary': data['base_salary']
        })
    
    return sorted(role_scores, key=lambda x: x['match_score'], reverse=True)

def recommend_courses(missing_skills):
    """Mock course recommendations."""
    recommendations = []
    for skill in missing_skills:
        recommendations.append(f"Advanced {skill.title()} Specialization on Coursera")
    return list(set(recommendations))[:5]

def get_project_recommendation(missing_skills):
    """Suggests a project."""
    return ["AI-Powered Career Intelligence System"] if not missing_skills else [f"Build a {missing_skills[0].title()} Application"]

def estimate_learning_time(missing_skills):
    return len(missing_skills) * 2

def get_interview_feedback(topic, user_answer, question_text=""):
    """
    Enhanced qualitative feedback analyzing professional tone, 
    STAR alignment, and impact-oriented language.
    """
    user_answer = user_answer.lower().strip()
    word_count = len(user_answer.split())
    
    feedback = []
    score = 0
    
    # 1. Length & Depth
    if word_count < 15:
        feedback.append("⚠️ **Conciseness vs. Detail:** Your answer is very brief. Try the **STAR method** (Situation, Task, Action, Result) to provide more context.")
        score += 10
    elif word_count < 50:
        feedback.append("✅ **Good Length:** You provided enough detail without being overly verbose.")
        score += 30
    else:
        feedback.append("⚠️ **Brevity Tip:** Your answer is quite long. Ensure you stay focused on the most impactful points.")
        score += 25

    # 2. Action-Oriented Language
    action_verbs = ['implemented', 'designed', 'developed', 'managed', 'led', 'analyzed', 'optimized', 'solved', 'built']
    found_actions = [v for v in action_verbs if v in user_answer]
    if found_actions:
        feedback.append(f"✅ **Strong Verbs:** Good use of active language ({', '.join(found_actions[:2])}).")
        score += 30
    else:
        feedback.append("⚠️ **Active Voice:** Use more action verbs like 'Implemented' or 'Optimized' to describe your contributions.")
        score += 10

    # 3. Result & Outcome Orientation
    outcome_keywords = ['result', 'achieved', 'improved', 'increased', 'decreased', 'saved', '%', 'impact', 'delivered', 'outcome']
    if any(w in user_answer for w in outcome_keywords) or re.search(r'\d+', user_answer):
        feedback.append("✅ **Impact Focused:** You mentioned clear outcomes or quantifiable results.")
        score += 40
    else:
        feedback.append("⚠️ **Result Orientation:** Focus more on the *impact* of your actions. Did it save time? Increase performance? Use numbers if possible.")
        score += 15

    # Final logic mapping
    if score >= 85:
        health = "Strong"
    elif score >= 50:
        health = "Competent"
    else:
        health = "Developing"

    return {
        'technical_health': health,
        'score': score
    }

def generate_targeted_bullets(missing_skills, target_role):
    return [f"Implemented technical solutions using {s} to improve system efficiency." for s in missing_skills[:3]]

def get_geographic_demand():
    data = {'City': ['Bangalore', 'Remote', 'Mumbai', 'Hyderabad', 'Pune'], 'Demand_Score': [95, 98, 85, 90, 80]}
    return pd.DataFrame(data)

def get_skill_proficiencies(user_skills, target_role):
    skills = ROLE_PATTERNS.get(target_role, {}).get('skills', [])[:5]
    data = []
    for s in skills:
        uv = 0.9 if s.lower() in [us.lower() for us in user_skills] else 0.2
        data.append({'Skill': s, 'Value': uv, 'Type': 'Your Level'})
        data.append({'Skill': s, 'Value': 1.0, 'Type': 'Requirement'})
    return pd.DataFrame(data)

def get_jobs_by_skill(skill):
    """Returns a list of roles associated with a specific skill."""
    skill = skill.lower()
    related_roles = []
    for role, data in ROLE_PATTERNS.items():
        if any(skill in s.lower() for s in data['skills']):
            related_roles.append(role)
    return related_roles if related_roles else ["General Software Role"]

def get_career_path_advice(role, matched_skills):
    """Provides personalized career advice based on the role and skills."""
    if not matched_skills:
        return f"To start a career as a {role}, focus on building a strong foundation in core technologies."
    
    advice = f"You have a great start for a {role} position with skills in {', '.join(matched_skills[:2])}. "
    advice += f"To advance, consider working on advanced projects and certifications specific to {role}."
    return advice
